import os
import sys
import json
import re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

def slugify(text):
    """Convert text to URL-friendly slug"""
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s-]', '', text)
    text = re.sub(r'[\s_]+', '-', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')

def extract_europe_package(file_path):
    """Extract package data from Europe Word file"""
    doc = Document(file_path)

    data = {
        'title': '',
        'duration': '',
        'description': '',
        'itinerary': [],
        'included': [],
        'notIncluded': [],
        'information': [],
        'hotels': {'threestar': [], 'fourstar': [], 'fivestar': []},
        'pricing': {
            'paxTiers': {},
            'childRates': {
                'age0to6': {'threestar': 0, 'fourstar': 0, 'fivestar': 0},
                'age6to12': {'threestar': 0, 'fourstar': 0, 'fivestar': 0}
            }
        },
        'destinations': '',
        'packageType': 'WITH_HOTEL',
        'tourType': 'PRIVATE',
        'transferType': 'PRIVATE',
        'region': 'Europe'
    }

    current_section = None
    current_day = None

    # Parse paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect title (usually centered, bold, first substantial text)
        if not data['title']:
            # Check for destination names
            cities = ['budapest', 'vienna', 'prague', 'zagreb', 'split', 'dubrovnik',
                     'helsinki', 'rovaniemi', 'finland', 'croatia', 'hungary', 'austria', 'czech']
            if any(city in text.lower() for city in cities):
                data['title'] = text
                continue

        # Detect duration
        if not data['duration'] and ('night' in text.lower() or 'gece' in text.lower()):
            if re.match(r'^\d+\s*(nights?|gece)\s*/\s*\d+\s*(days?|gün)', text, re.I):
                data['duration'] = text
                continue

        # Detect day entries - multiple formats
        # Format 1: Day 1 - Title (B/L/D)
        day_match = re.match(r'^Day\s*(\d+)\s*[-–]\s*(.+?)\s*\(([^)]*)\)', text, re.I)
        if not day_match:
            # Format 2: Day 1 - Title    (B/L/D) with tabs
            day_match = re.match(r'^Day\s*(\d+)\s*[-–]\s*(.+?)\s+\(([^)]*)\)', text, re.I)
        if not day_match:
            # Format 3: 1. Gün - Title (K/Ö/A) Turkish
            day_match = re.match(r'^(\d+)\.\s*(?:Gün|Day)\s*[-–]\s*(.+?)\s*\(([^)]*)\)', text, re.I)

        if day_match:
            current_day = {
                'day': int(day_match.group(1)),
                'title': day_match.group(2).strip(),
                'description': '',
                'meals': day_match.group(3).strip() or '-'
            }
            data['itinerary'].append(current_day)
            current_section = 'itinerary'
            continue

        # Continue day description
        if current_section == 'itinerary' and current_day:
            if text.startswith('***') or text.lower().startswith('end of service'):
                current_section = None
                current_day = None
            elif not re.match(r'^(Day\s*\d+|\d+\.\s*(?:Gün|Day))', text, re.I):
                if text.lower() not in ['inclusions', 'exclusions', 'information', 'hotel options',
                                        'package rates', 'dahil olan hizmetler', 'dahil olmayan hizmetler']:
                    if current_day['description']:
                        current_day['description'] += ' '
                    current_day['description'] += text
                    continue

        # Detect sections
        text_lower = text.lower()
        if text_lower in ['inclusions', 'dahil olan hizmetler']:
            current_section = 'included'
            continue
        elif text_lower in ['exclusions', 'dahil olmayan hizmetler']:
            current_section = 'notIncluded'
            continue
        elif text_lower in ['information', 'important information', 'bilgilendirme', 'önemli bilgiler']:
            current_section = 'information'
            continue
        elif text_lower in ['hotel options', 'otel seçenekleri', 'hotel information']:
            current_section = 'hotels'
            continue
        elif text_lower in ['package rates', 'paket fiyatları']:
            current_section = 'pricing'
            continue

        # Add to current section
        if current_section == 'included' and text:
            clean = text.rstrip(',').rstrip('.')
            if clean and clean.lower() not in ['exclusions', 'dahil olmayan hizmetler']:
                data['included'].append(clean)
        elif current_section == 'notIncluded' and text:
            clean = text.rstrip(',').rstrip('.')
            if clean and clean.lower() not in ['information', 'hotel options', 'bilgilendirme']:
                data['notIncluded'].append(clean)
        elif current_section == 'information' and text:
            if text.lower() not in ['hotel options', 'package rates', 'otel seçenekleri']:
                data['information'].append(text)

    # Parse tables
    for table in doc.tables:
        if len(table.rows) < 2:
            continue

        first_cell = table.rows[0].cells[0].text.strip().lower()

        # Hotels table
        if any(x in first_cell for x in ['city', 'şehir', 'location', 'lokasyon']):
            for row in table.rows[1:]:
                if len(row.cells) >= 4:
                    city = row.cells[0].text.strip()
                    h3 = row.cells[1].text.strip().replace('\n', ' ')
                    h4 = row.cells[2].text.strip().replace('\n', ' ')
                    h5 = row.cells[3].text.strip().replace('\n', ' ')
                    if city and h3:
                        data['hotels']['threestar'].append(f"{city}: {h3}")
                    if city and h4:
                        data['hotels']['fourstar'].append(f"{city}: {h4}")
                    if city and h5:
                        data['hotels']['fivestar'].append(f"{city}: {h5}")

        # Pricing table
        elif any(x in first_cell for x in ['pax', 'kişi', 'person']):
            for row in table.rows[1:]:
                if len(row.cells) >= 4:
                    label = row.cells[0].text.strip().upper()

                    def extract_price(cell_text):
                        match = re.search(r'(\d+)', cell_text.replace(',', '').replace('.', ''))
                        return int(match.group(1)) if match else 0

                    p3 = extract_price(row.cells[1].text)
                    p4 = extract_price(row.cells[2].text)
                    p5 = extract_price(row.cells[3].text)

                    # Map to pax tiers
                    if '10' in label and ('PAX' in label or 'KİŞİ' in label):
                        data['pricing']['paxTiers']['10'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '8' in label and ('PAX' in label or 'KİŞİ' in label):
                        data['pricing']['paxTiers']['8'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '6' in label and ('PAX' in label or 'KİŞİ' in label):
                        data['pricing']['paxTiers']['6'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '4' in label and ('PAX' in label or 'KİŞİ' in label):
                        data['pricing']['paxTiers']['4'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '2' in label and ('PAX' in label or 'KİŞİ' in label):
                        data['pricing']['paxTiers']['2'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif 'SINGLE' in label or 'TEK' in label:
                        for tier in data['pricing']['paxTiers']:
                            data['pricing']['paxTiers'][tier]['threestar']['singleSupplement'] = p3
                            data['pricing']['paxTiers'][tier]['fourstar']['singleSupplement'] = p4
                            data['pricing']['paxTiers'][tier]['fivestar']['singleSupplement'] = p5
                    elif 'CHILD' in label or 'ÇOCUK' in label:
                        if '0-6' in label or '0' in label.split()[0]:
                            data['pricing']['childRates']['age0to6'] = {'threestar': p3, 'fourstar': p4, 'fivestar': p5}
                        elif '6-12' in label or '12' in label:
                            data['pricing']['childRates']['age6to12'] = {'threestar': p3, 'fourstar': p4, 'fivestar': p5}

    # Extract destinations from title
    title_lower = data['title'].lower()
    destinations = []

    city_map = {
        'budapest': 'Budapest',
        'vienna': 'Vienna',
        'prague': 'Prague',
        'zagreb': 'Zagreb',
        'split': 'Split',
        'dubrovnik': 'Dubrovnik',
        'helsinki': 'Helsinki',
        'rovaniemi': 'Rovaniemi',
        'hvar': 'Hvar'
    }

    for key, city in city_map.items():
        if key in title_lower:
            destinations.append(city)

    # Also check itinerary
    for day in data['itinerary']:
        for key, city in city_map.items():
            if key in day['title'].lower() and city not in destinations:
                destinations.append(city)

    data['destinations'] = ', '.join(destinations) if destinations else 'Europe'

    # Create description
    if data['itinerary']:
        data['description'] = f"Discover the beauty of {data['destinations']} with this {data['duration']} package. "
        if data['itinerary'][0]['description']:
            data['description'] += data['itinerary'][0]['description'][:200] + "..."

    # Generate highlights
    highlights = []
    for day in data['itinerary']:
        if any(x in day['title'].lower() for x in ['tour', 'visit', 'explore', 'cruise']):
            highlights.append(day['title'])
    data['highlights'] = highlights[:5] if highlights else [
        f"Explore {data['destinations']}",
        "Professional English-speaking guide",
        "Private airport transfers",
        "Quality 4-star hotel accommodations"
    ]

    return data


def process_europe_packages():
    """Process all Europe Word files"""
    europe_path = r"C:\Users\fatih\Desktop\2026 packages\EUROPE 2026"

    all_packages = []

    print(f"\n=== Processing Europe Packages ===")
    print(f"Path: {europe_path}\n")

    # Only process .docx files (not Turkish translations or temp files)
    for filename in sorted(os.listdir(europe_path)):
        if not filename.endswith('.docx'):
            continue
        if filename.startswith('~$'):
            continue
        if 'TR' in filename or 'TURKCE' in filename:
            continue  # Skip Turkish versions for now

        file_path = os.path.join(europe_path, filename)

        try:
            pkg_data = extract_europe_package(file_path)

            # Extract package number from filename
            num_match = re.match(r'^(\d+)', filename)
            pkg_num = num_match.group(1) if num_match else '00'

            # Create unique package ID for Europe
            pkg_data['packageId'] = f"EUR-PVT-{pkg_num}"
            pkg_data['slug'] = slugify(f"{pkg_data['title']}-private")
            pkg_data['region'] = 'Europe'  # Important: Set region for filtering

            # Set default image based on destination
            first_dest = pkg_data['destinations'].split(',')[0].strip().lower() if pkg_data['destinations'] else 'europe'
            pkg_data['image'] = f"/images/packages/{first_dest}-tour.jpg"

            all_packages.append(pkg_data)
            print(f"  ✓ Extracted: {pkg_data['packageId']} - {pkg_data['title']}")

        except Exception as e:
            print(f"  ✗ Error: {filename} - {e}")

    return all_packages


def main():
    """Main function to extract and merge packages"""

    # Extract Europe packages
    europe_packages = process_europe_packages()

    # Load existing Turkey packages
    data_path = os.path.join(os.path.dirname(__file__), '..', 'data', '2026-packages.json')

    if os.path.exists(data_path):
        with open(data_path, 'r', encoding='utf-8') as f:
            existing_packages = json.load(f)
        print(f"\nLoaded {len(existing_packages)} existing Turkey packages")
    else:
        existing_packages = []
        print("\nNo existing packages found")

    # Remove any existing Europe packages (to avoid duplicates)
    turkey_only = [p for p in existing_packages if not p.get('packageId', '').startswith('EUR-')]

    # Ensure all Turkey packages have region set
    for pkg in turkey_only:
        if 'region' not in pkg:
            pkg['region'] = 'Turkey'

    print(f"Turkey packages: {len(turkey_only)}")

    # Merge packages
    all_packages = turkey_only + europe_packages

    # Save merged JSON
    with open(data_path, 'w', encoding='utf-8') as f:
        json.dump(all_packages, f, indent=2, ensure_ascii=False)

    print(f"\n=== COMPLETE ===")
    print(f"Total packages: {len(all_packages)}")
    print(f"  - Turkey: {len(turkey_only)}")
    print(f"  - Europe: {len(europe_packages)}")
    print(f"Saved to: {data_path}")

    # Also save Europe-only JSON for reference
    europe_only_path = os.path.join(os.path.dirname(__file__), '..', 'data', 'europe-packages.json')
    with open(europe_only_path, 'w', encoding='utf-8') as f:
        json.dump(europe_packages, f, indent=2, ensure_ascii=False)
    print(f"Europe packages also saved to: {europe_only_path}")


if __name__ == '__main__':
    main()
