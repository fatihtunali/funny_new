import os
import sys
import json
import re
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

# Spanish translation dictionary for common tourism terms
TRANSLATIONS = {
    # Tour types
    'Group Tour': 'Tour en Grupo',
    'Private Tour': 'Tour Privado',
    'Land Only': 'Solo Terrestre',
    'Nights': 'Noches',

    # Meals
    'Breakfast': 'Desayuno',
    'Lunch': 'Almuerzo',
    'Dinner': 'Cena',
    'B': 'D',  # Breakfast -> Desayuno
    'L': 'A',  # Lunch -> Almuerzo
    'D': 'C',  # Dinner -> Cena

    # Common terms
    'Day': 'Día',
    'Arrival': 'Llegada',
    'Departure': 'Salida',
    'Transfer': 'Traslado',
    'Tour': 'Tour',
    'Full Day': 'Día Completo',
    'Half Day': 'Medio Día',
    'Free Day': 'Día Libre',
    'at leisure': 'tiempo libre',
    'hotel': 'hotel',
    'airport': 'aeropuerto',
    'cruise': 'crucero',
    'visit': 'visitar',
    'explore': 'explorar',

    # Cities
    'Istanbul': 'Estambul',
    'Cappadocia': 'Capadocia',
    'Kusadasi': 'Kusadasi',
    'Antalya': 'Antalya',
    'Ephesus': 'Éfeso',
    'Pamukkale': 'Pamukkale',

    # Inclusions
    'nights accommodation': 'noches de alojamiento',
    'Meals as per itinerary': 'Comidas según itinerario',
    'Return airport transfers': 'Traslados de ida y vuelta al aeropuerto',
    'Private basis': 'base privada',
    'Professional English-Speaking Guidance': 'Guía profesional de habla inglesa',
    'Group Tours': 'Tours en grupo',
    'entrance fees': 'entradas incluidas',
    'Local Taxes': 'Impuestos locales',

    # Exclusions
    'Flights': 'Vuelos',
    'Personal expenses': 'Gastos personales',
    'Drinks at meals': 'Bebidas en las comidas',
    'Tips': 'Propinas',
    'porterage': 'maleteros',
    'driver': 'conductor',
    'guide': 'guía',
}

def translate_to_spanish(text):
    """Simple translation using dictionary"""
    result = text
    for en, es in TRANSLATIONS.items():
        result = re.sub(re.escape(en), es, result, flags=re.IGNORECASE)
    return result

def translate_title(title):
    """Translate package title to Spanish"""
    result = title
    result = result.replace('Nights Group Tour', 'Noches Tour en Grupo')
    result = result.replace('Nights Private Tour', 'Noches Tour Privado')
    result = result.replace(' - Land Only', ' - Solo Terrestre')
    result = result.replace('Istanbul', 'Estambul')
    result = result.replace('Cappadocia', 'Capadocia')
    result = result.replace('Ephesus', 'Éfeso')
    return result

def translate_itinerary(itinerary):
    """Translate itinerary to Spanish"""
    translated = []
    for day in itinerary:
        es_day = {
            'day': day['day'],
            'title': translate_day_title(day['title']),
            'description': translate_description(day['description']),
            'meals': translate_meals(day['meals'])
        }
        translated.append(es_day)
    return translated

def translate_day_title(title):
    """Translate day title"""
    result = title
    result = result.replace('Fly', 'Vuelo')
    result = result.replace('Full Day Tour', 'Tour de Día Completo')
    result = result.replace('Half Day Tour', 'Tour de Medio Día')
    result = result.replace('Free Day', 'Día Libre')
    result = result.replace('Arrival', 'Llegada')
    result = result.replace('Departure', 'Salida')
    result = result.replace('Istanbul', 'Estambul')
    result = result.replace('Cappadocia', 'Capadocia')
    result = result.replace('Ephesus', 'Éfeso')
    result = result.replace('Bosphorus Cruise', 'Crucero por el Bósforo')
    result = result.replace('incl.', 'incl.')
    return result

def translate_description(desc):
    """Translate description to Spanish"""
    if not desc:
        return ''
    result = desc
    # Common phrases
    result = result.replace('After breakfast at the hotel', 'Después del desayuno en el hotel')
    result = result.replace('After breakfast', 'Después del desayuno')
    result = result.replace('Arrive to', 'Llegada a')
    result = result.replace('Arrival transfer to the hotel', 'Traslado de llegada al hotel')
    result = result.replace('check-in', 'registro')
    result = result.replace('The rest of the day is free', 'El resto del día es libre')
    result = result.replace('at leisure', 'tiempo libre')
    result = result.replace('guided tour', 'tour guiado')
    result = result.replace('your tour will start', 'su tour comenzará')
    result = result.replace('departure airport transfer', 'traslado al aeropuerto de salida')
    result = result.replace('overnight', 'pernocte')
    result = result.replace('Overnight', 'Pernocte')
    # Places
    result = result.replace('Istanbul', 'Estambul')
    result = result.replace('Cappadocia', 'Capadocia')
    result = result.replace('Ephesus', 'Éfeso')
    result = result.replace('Blue Mosque', 'Mezquita Azul')
    result = result.replace('Topkapi Palace', 'Palacio de Topkapi')
    result = result.replace('Grand Bazaar', 'Gran Bazar')
    result = result.replace('Spice Bazaar', 'Bazar de las Especias')
    result = result.replace('Bosphorus', 'Bósforo')
    result = result.replace('St.Sophia', 'Santa Sofía')
    result = result.replace('Hagia Sophia', 'Santa Sofía')
    result = result.replace('Hippodrome', 'Hipódromo')
    return result

def translate_meals(meals):
    """Translate meal codes"""
    if not meals or meals == '-':
        return '-'
    result = meals
    result = result.replace('B', 'D')  # Breakfast -> Desayuno
    result = result.replace('L', 'A')  # Lunch -> Almuerzo
    # Note: D for Dinner stays as C (Cena) but we need to be careful
    # Only replace D that's for Dinner, not Desayuno we just added
    return result

def translate_list(items):
    """Translate list of items"""
    translated = []
    for item in items:
        translated.append(translate_to_spanish(item))
    return translated

def slugify(text):
    """Convert text to URL-friendly slug"""
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s-]', '', text)
    text = re.sub(r'[\s_]+', '-', text)
    text = re.sub(r'-+', '-', text)
    return text.strip('-')

def create_better_title(destinations, duration, tour_type, package_type):
    """Create a descriptive package title"""
    # Extract nights from duration
    nights_match = re.search(r'(\d+)\s*nights?', duration, re.I)
    nights = nights_match.group(1) if nights_match else ''

    # Format destinations
    dest_list = [d.strip() for d in destinations.split(',')]
    if len(dest_list) == 1:
        dest_str = dest_list[0]
    elif len(dest_list) == 2:
        dest_str = f"{dest_list[0]} & {dest_list[1]}"
    else:
        dest_str = f"{dest_list[0]}, {', '.join(dest_list[1:-1])} & {dest_list[-1]}"

    # Tour type label
    tour_label = "Group Tour" if tour_type == "SIC" else "Private Tour"

    # Package type suffix
    if package_type == "LAND_ONLY":
        type_suffix = " - Land Only"
    else:
        type_suffix = ""

    # Create title
    if nights:
        return f"{dest_str} {nights} Nights {tour_label}{type_suffix}"
    else:
        return f"{dest_str} {tour_label}{type_suffix}"

def extract_package_from_word(file_path, package_type='WITH_HOTEL', tour_type='SIC'):
    """Extract all package data from a Word file"""
    doc = Document(file_path)

    # Initialize data structure
    data = {
        'title': '',
        'duration': '',
        'description': '',
        'itinerary': [],
        'included': [],
        'notIncluded': [],
        'information': [],
        'hotels': {'threestar': [], 'fourstar': [], 'fivestar': []},
        'pricing': {
            'paxTiers': {},
            'childRates': {
                'age0to6': {'threestar': 0, 'fourstar': 0, 'fivestar': 0},
                'age6to12': {'threestar': 0, 'fourstar': 0, 'fivestar': 0}
            }
        },
        'destinations': '',
        'packageType': package_type,
        'tourType': tour_type,
        'transferType': 'PRIVATE'  # All packages use private transfers
    }

    current_section = None
    current_day = None

    # Parse paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Detect title (first significant paragraph, usually bold)
        if not data['title'] and 'package' in text.lower():
            data['title'] = text
            continue

        # Detect duration
        if not data['duration'] and ('night' in text.lower() or 'day' in text.lower()):
            if re.match(r'^\d+\s*nights?\s*/\s*\d+\s*days?', text, re.I):
                data['duration'] = text
                continue

        # Detect day entries
        day_match = re.match(r'^Day\s*(\d+)\s*[-–]\s*(.+?)\s*\(([^)]*)\)', text, re.I)
        if day_match:
            current_day = {
                'day': int(day_match.group(1)),
                'title': day_match.group(2).strip(),
                'description': '',
                'meals': day_match.group(3).strip() or '-'
            }
            data['itinerary'].append(current_day)
            current_section = 'itinerary'
            continue

        # Continue day description
        if current_section == 'itinerary' and current_day and not text.startswith('Day '):
            if text.startswith('***'):
                current_section = None
                current_day = None
            elif text.lower() not in ['inclusions', 'exclusions', 'information', 'hotel options', 'package rates']:
                if current_day['description']:
                    current_day['description'] += ' '
                current_day['description'] += text
                continue

        # Detect sections
        if text.lower() == 'inclusions':
            current_section = 'included'
            continue
        elif text.lower() == 'exclusions':
            current_section = 'notIncluded'
            continue
        elif text.lower() == 'information':
            current_section = 'information'
            continue

        # Add to current section
        if current_section == 'included' and text and not text.startswith('***'):
            # Clean up the text
            clean_text = text.rstrip(',').rstrip('.')
            if clean_text and clean_text.lower() not in ['exclusions', 'information']:
                data['included'].append(clean_text)
        elif current_section == 'notIncluded' and text and not text.startswith('***'):
            clean_text = text.rstrip(',').rstrip('.')
            if clean_text and clean_text.lower() not in ['information', 'hotel options']:
                data['notIncluded'].append(clean_text)
        elif current_section == 'information' and text and not text.startswith('***'):
            if text.lower() not in ['hotel options', 'package rates']:
                data['information'].append(text)

    # Parse tables
    for table in doc.tables:
        first_cell = table.rows[0].cells[0].text.strip().lower()

        # Hotels table
        if 'city' in first_cell:
            for row in table.rows[1:]:
                if len(row.cells) >= 4:
                    city = row.cells[0].text.strip()
                    h3 = row.cells[1].text.strip().replace('\n', ' ').replace('  ', ' ')
                    h4 = row.cells[2].text.strip().replace('\n', ' ').replace('  ', ' ')
                    h5 = row.cells[3].text.strip().replace('\n', ' ').replace('  ', ' ')
                    if city:
                        if h3:
                            data['hotels']['threestar'].append(f"{city}: {h3}")
                        if h4:
                            data['hotels']['fourstar'].append(f"{city}: {h4}")
                        if h5:
                            data['hotels']['fivestar'].append(f"{city}: {h5}")

        # Pricing table
        elif 'pax' in first_cell or 'dbl' in first_cell:
            for row in table.rows[1:]:
                if len(row.cells) >= 4:
                    label = row.cells[0].text.strip().upper()

                    # Extract numeric values
                    def extract_price(cell_text):
                        match = re.search(r'(\d+)', cell_text.replace(',', ''))
                        return int(match.group(1)) if match else 0

                    p3 = extract_price(row.cells[1].text)
                    p4 = extract_price(row.cells[2].text)
                    p5 = extract_price(row.cells[3].text)

                    # Map to pax tiers
                    if '10' in label and 'PAX' in label:
                        data['pricing']['paxTiers']['10'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '8' in label and 'PAX' in label:
                        data['pricing']['paxTiers']['8'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '6' in label and 'PAX' in label:
                        data['pricing']['paxTiers']['6'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '4' in label and 'PAX' in label:
                        data['pricing']['paxTiers']['4'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif '2' in label and 'PAX' in label:
                        data['pricing']['paxTiers']['2'] = {
                            'threestar': {'double': p3, 'singleSupplement': 0},
                            'fourstar': {'double': p4, 'singleSupplement': 0},
                            'fivestar': {'double': p5, 'singleSupplement': 0}
                        }
                    elif 'SINGLE' in label:
                        # Add single supplement to all tiers
                        for tier in data['pricing']['paxTiers']:
                            data['pricing']['paxTiers'][tier]['threestar']['singleSupplement'] = p3
                            data['pricing']['paxTiers'][tier]['fourstar']['singleSupplement'] = p4
                            data['pricing']['paxTiers'][tier]['fivestar']['singleSupplement'] = p5
                    elif 'CHILD' in label and ('0-6' in label or '0' in label.split()[0] if len(label.split()) > 0 else False):
                        data['pricing']['childRates']['age0to6'] = {'threestar': p3, 'fourstar': p4, 'fivestar': p5}
                    elif 'CHILD' in label and ('6-12' in label or '12' in label):
                        data['pricing']['childRates']['age6to12'] = {'threestar': p3, 'fourstar': p4, 'fivestar': p5}

    # Extract destinations from title or itinerary
    destinations = []
    title_lower = data['title'].lower()
    cities = ['istanbul', 'cappadocia', 'kusadasi', 'antalya', 'ephesus', 'pamukkale', 'bodrum', 'fethiye']
    for city in cities:
        if city in title_lower:
            destinations.append(city.title())

    # Also check itinerary for destinations
    for day in data['itinerary']:
        for city in cities:
            if city in day['title'].lower() and city.title() not in destinations:
                destinations.append(city.title())

    data['destinations'] = ', '.join(destinations) if destinations else 'Turkey'

    # Create better title
    data['title'] = create_better_title(data['destinations'], data['duration'], tour_type, package_type)

    # Create description from first day or general intro
    if data['itinerary']:
        data['description'] = f"Experience the best of {data['destinations']} with this {data['duration']} package. "
        if data['itinerary'][0]['description']:
            data['description'] += data['itinerary'][0]['description'][:200] + "..."

    # Add Spanish translations
    data['titleEs'] = translate_title(data['title'])
    data['descriptionEs'] = f"Experimente lo mejor de {translate_to_spanish(data['destinations'])} con este paquete de {data['duration'].replace('Nights', 'Noches').replace('Days', 'Días')}. "
    if data['itinerary'] and data['itinerary'][0]['description']:
        data['descriptionEs'] += translate_description(data['itinerary'][0]['description'][:200]) + "..."
    data['itineraryEs'] = translate_itinerary(data['itinerary'])
    data['includedEs'] = translate_list(data['included'])
    data['notIncludedEs'] = translate_list(data['notIncluded'])
    data['highlightsEs'] = translate_list(data.get('highlights', []))

    return data


def process_all_packages():
    """Process all Word files and create JSON data"""
    base_path = r"C:\Users\fatih\Desktop\2026 packages\website prices"

    folders = [
        ("2026 PACKAGE WITH SIC TOURS & PVT AIRPORT TRANSFERS - FT - done", "WITH_HOTEL", "SIC"),
        ("2026 PACKAGES WITH PVT TOURS & PVT AIRPORT TRANSFERS - FT", "WITH_HOTEL", "PRIVATE"),
        ("ONLY LAND SIC TOURS & PVT AIRPORT TRANSFERS - FT", "LAND_ONLY", "SIC"),
        ("ONLY LAND PVT TOURS & PVT AIRPORT TRANSFERS - FT", "LAND_ONLY", "PRIVATE"),
    ]

    all_packages = []

    for folder, package_type, tour_type in folders:
        folder_path = os.path.join(base_path, folder)
        if not os.path.exists(folder_path):
            continue

        print(f"\n=== Processing: {folder} ===")

        for filename in sorted(os.listdir(folder_path)):
            if not filename.endswith('.docx') or filename.startswith('~$'):
                continue

            file_path = os.path.join(folder_path, filename)

            try:
                pkg_data = extract_package_from_word(file_path, package_type, tour_type)

                # Extract package number from filename
                num_match = re.match(r'^(\d+)', filename)
                pkg_num = num_match.group(1) if num_match else '00'

                # Create unique package ID based on type
                type_prefix = {
                    ('WITH_HOTEL', 'SIC'): 'SIC',
                    ('WITH_HOTEL', 'PRIVATE'): 'PVT',
                    ('LAND_ONLY', 'SIC'): 'LAND-SIC',
                    ('LAND_ONLY', 'PRIVATE'): 'LAND-PVT',
                }
                prefix = type_prefix.get((package_type, tour_type), 'PKG')

                pkg_data['packageId'] = f"{prefix}-{pkg_num}"
                pkg_data['slug'] = slugify(f"{pkg_data['title']}-{tour_type.lower()}")

                # Set default image
                pkg_data['image'] = f"/images/packages/{pkg_data['slug']}.jpg"

                # Generate highlights from itinerary
                highlights = []
                for day in pkg_data['itinerary']:
                    if 'tour' in day['title'].lower() or 'visit' in day['title'].lower():
                        highlights.append(day['title'])
                pkg_data['highlights'] = highlights[:5] if highlights else [
                    f"Explore {pkg_data['destinations']}",
                    "Professional English-speaking guide",
                    "All transfers included",
                    "Quality hotel accommodations"
                ]

                all_packages.append(pkg_data)
                print(f"  Extracted: {pkg_data['packageId']} - {pkg_data['title']}")

            except Exception as e:
                print(f"  Error: {filename} - {e}")

    # Save to JSON
    output_dir = os.path.join(os.path.dirname(__file__), '..', 'data')
    os.makedirs(output_dir, exist_ok=True)

    output_path = os.path.join(output_dir, '2026-packages.json')
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(all_packages, f, indent=2, ensure_ascii=False)

    print(f"\n=== COMPLETE ===")
    print(f"Extracted {len(all_packages)} packages")
    print(f"Saved to: {output_path}")

    return all_packages


if __name__ == '__main__':
    process_all_packages()
